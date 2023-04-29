import csv
from msilib.schema import ComboBox
from pickle import GLOBAL
import string
from tkinter import *
from tkinter import ttk
from tkinter.scrolledtext import ScrolledText
import random
from types import NoneType
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from tkinter import messagebox
import os
from os import listdir
from os.path import isfile, join
from datetime import datetime
import math
import time


all_questions = []
chapterIndex = 5
COURSE_NAME = "ICS4U"
TEST_NAME = ""
TEST_WEIGHT = ""
TEST_TIME_LIMIT = ""
FILE_NAME = ""
FIRST_FILE = ""
COPIES = 1
#pulls all questions from the csv file and stores as a 2d array
with open('files.csv') as csv_file:
    csv_reader = csv.reader(csv_file, delimiter=",")
    for row in csv_reader:
        all_questions.append(row)

def chapters_reader(questions,index):
    all_chaps = []
    for question in questions:
        try:
            testing = float(question[index])
            all_chaps.append(question[index])
            
        except:
            ValueError
            
    all_chaps = list(dict.fromkeys(all_chaps))
    for i in range(len(all_chaps)):
        all_chaps[i] = str(all_chaps[i])
    return all_chaps

all_chapters = chapters_reader(all_questions,chapterIndex)


#GUI for the program 
root = Tk()
root.geometry("800x600")
root.title("Quiz Generator")

explanationLabel = Label(root,text=" In the first entry box in each row, list how many questions you want to get. The second and third entry boxes determine which chapters\n the questions are determined from. For example, entering 10 in the first box, 1.1 in the second box, and 1.5 in the third box would\n make the program choose 10 random questions from 1.1,1.2,1.3,1.4, and 1.5. Leave rows empty if you do not need to use them.", justify=LEFT)


#this function limits the third column options based on the second column choice for the first row.
def option_reduction1(options):
    options = str(s1_1.get())
    arr1 = options.split(".")
    newArray = []
    for i in range(len(all_chapters)):
        arr2 = all_chapters[i].split(".")
        if arr1[0] == arr2[0] and arr1[1] <= arr2[1]:
            newArray.append(all_chapters[i])
    s1_2.config(value = newArray)
    s1_2.current(0)

#this function limits the third column options based on the second column choice for the second row.
def option_reduction2(options):
    options = str(s2_1.get())
    arr1 = options.split(".")
    newArray = []
    for i in range(len(all_chapters)):
        arr2 = all_chapters[i].split(".")
        if arr1[0] == arr2[0] and arr1[1] <= arr2[1]:
            newArray.append(all_chapters[i])
    s2_2.config(value = newArray)
    s2_2.current(0)

#this function limits the third column options based on the second column choice for the third row.
def option_reduction3(options):
    options = str(s3_1.get())
    arr1 = options.split(".")
    newArray = []
    for i in range(len(all_chapters)):
        arr2 = all_chapters[i].split(".")
        if arr1[0] == arr2[0] and arr1[1] <= arr2[1]:
            newArray.append(all_chapters[i])
    s3_2.config(value = newArray)
    s3_2.current(0)

#this function limits the third column options based on the second column choice for the fourth row.
def option_reduction4(options):
    options = str(s4_1.get())
    arr1 = options.split(".")
    newArray = []
    for i in range(len(all_chapters)):
        arr2 = all_chapters[i].split(".")
        if arr1[0] == arr2[0] and arr1[1] <= arr2[1]:
            newArray.append(all_chapters[i])
    s4_2.config(value = newArray)
    s4_2.current(0)

#this function limits the third column options based on the second column choice for the fifth row.
def option_reduction5(options):
    options = str(s5_1.get())
    arr1 = options.split(".")
    newArray = []
    for i in range(len(all_chapters)):
        arr2 = all_chapters[i].split(".")
        if arr1[0] == arr2[0] and arr1[1] <= arr2[1]:
            newArray.append(all_chapters[i])
    s5_2.config(value = newArray)
    s5_2.current(0)

#uses the entries given to make a pdf of random questions
stored_ranges = []
stored_amnts = []
stored_toggles = []

def closeWindow(window,final):
    if final:
        root.destroy()
        os.chdir("./Quizes/")
        os.system("start " + FILE_NAME + "")
        
    window.destroy()
    

def checkInsides(questions,s1,s2):
    errorText = ""
    if questions == questions1:
        errorText = "The first"
    elif questions == questions2:
        errorText = "The second"
    elif questions == questions3:
        errorText = "The third"
    elif questions == questions4:
        errorText = "The fourth"
    elif questions == questions5:
        errorText = "The fifth"
    if (questions.get() == "" and s1.get() == "" and s2.get() == ""):
        pass
    elif (questions.get() == "" and s1.get() != "" and s2.get() != ""):
        return [False,errorText]
    elif (questions.get() != "" and (s1.get() == "" or s2.get() == "")):
        return [False,errorText]
    else:
        return True

def checkZero(value):
    if value == "":
        return True
    else:
        return False


def getValues(box1,box2,box3,box4,window):
    global COURSE_NAME
    global TEST_NAME
    global TEST_TIME_LIMIT
    global TEST_WEIGHT 
    global COPIES
 #   COURSE_NAME = box1.get()

    TEST_NAME = box1.get()
    TEST_WEIGHT = box2.get()
    TEST_TIME_LIMIT = box3.get()
    COPIES = box4.get()
    root.deiconify()
    window.destroy()

def on_closing():
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        root.destroy()
    

def getCoverInfo():
    root.withdraw()
    win = Toplevel(root)
    win.geometry("460x360")
    windowWidth = win.winfo_reqwidth()
    windowHeight = win.winfo_reqheight()
 
    positionRight = int(win.winfo_screenwidth()/2 - windowWidth/2)
    positionDown = int(win.winfo_screenheight()/2 - windowHeight/2)
    win.geometry("+{}+{}".format(positionRight-175, positionDown-100))
    values_returned = []
    label1 = Label(win,text="ICS4U",font=("Arial",24))
    label2 = Label(win,text="Enter the amount of time in minutes for this quiz")
    label3 = Label(win,text="Enter the weight of the quiz in percent here")
    label4 = Label(win,text="Enter the name of the quiz/test")
    label5 = Label(win,text="How many copies of the quiz will be made?")

    minutes_box = Entry(win)
    minutes_box.insert(END, "10")
    percentage_box = Entry(win)
    percentage_box.insert(END, "3%")
    quiznum_box = Entry(win)
    quiznum_box.insert(END, "Quiz 4A")
    copies_variable = StringVar()

    keepvalue = copies_variable.get()
    copies_options = ["1","2","3","4","5","6","7","8","9","10"]
    
    copies_box = ttk.Combobox(win, textvariable=keepvalue, value = copies_options , width=15)
    copies_box.current(0)


    label1.grid(row=0,column=0,padx=10,pady=20,columnspan=2)
    label2.grid(row=1,column=0,padx=10,pady=20)
    label3.grid(row=2,column=0,padx=10,pady=20)
    label4.grid(row=3,column=0,padx=10,pady=20)
    label5.grid(row=4,column=0,padx=10,pady=20)
    minutes_box.grid(row=1,column=1,padx=10,pady=20)
    percentage_box.grid(row=2,column=1,padx=10,pady=20)
    quiznum_box.grid(row=3,column=1,padx=10,pady=20)
    copies_box.grid(row=4,column=1,padx=10,pady=20)

    ok_button = Button(win,text="submit",command= lambda:getValues(quiznum_box,minutes_box,percentage_box,copies_box,win))
    ok_button.grid(row=5,column=0,columnspan=2)
    win.protocol("WM_DELETE_WINDOW", on_closing)






def coverTitle(quiz,courseName,quiznum,minutes,percentage,questions):
    ics = quiz.add_paragraph()
    icsstyles = quiz.styles
    try:
        font_charstyle = icsstyles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    except:
        ValueError
        font_charstyle = quiz.styles["CommentsStyle"]
    
    font_object = font_charstyle.font
    font_object.size = Pt(20)
    font_object.name = 'Times New Roman'
    ics.add_run("\n"+courseName+"\n"+quiznum,style="CommentsStyle").bold=True
    ics.alignment = 1 
    ics.paragraph_format.line_spacing = Inches(0.6)
    cover_title = quiz.add_paragraph(percentage+" of Final Grade\n"+minutes+" minutes\n"+questions+" Marks\nClass Section __________________\nFirst Name   __________________\nLast Name   __________________\nStudent #   __________________\nDate      __________________")    
    run = cover_title.add_run("\nInstructions")
    run.underline = True
    run.bold = True
    run1 = cover_title.add_run("\nThis is a closed book test.\nNo calculator or electronic aids are allowed.\nAnswer all questions.\nMark your answers on the GradeMaster in ")
    pencil = cover_title.add_run("pencil").bold = True
    run2 = cover_title.add_run("\nYou do not have to put your student number on the GradeMaster.\nMake sure you write your name on both this Quiz and the GradeMaster.")
    run3 = cover_title.add_run("\nTo receive a grade, submit this Quiz and the GradeMaster.").bold = True
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.line_spacing = Inches(0.4)
    cover_titleFont = cover_title.style.font
    cover_titleFont.name = "Times New Roman"
    cover_titleFont.size = Pt(11)

def printCheck():
    print(stored_amnts)
    print(stored_ranges)
    print(stored_toggles)

def makePDF():
    global TEST_NAME
    global FILE_NAME
    global COPIES
    global FIRST_FILE
    global stored_amnts
    alphabet = "abcdefghijklmnopqrstuvwxyz"
    printCheck()
    for j in range(int(COPIES)):

        h = str(j+1)
        questionsReturned = []
        questionIndexes = []
        row1 = checkInsides(questions1,s1_1,s1_2)
        row2 = checkInsides(questions2,s2_1,s2_2)
        row3 = checkInsides(questions3,s3_1,s3_2)
        row4 = checkInsides(questions4,s4_1,s4_2)
        row5 = checkInsides(questions5,s5_1,s5_2)
        allErrors = []
        if type(row1) != NoneType and type(row1) == list:
            if row1[0] == False:
                allErrors.append("first")
        if type(row2) != NoneType and type(row2) == list:
            if row2[0] == False:
                allErrors.append("second")
        if type(row3) != NoneType and type(row3) == list:
            if row3[0] == False:
                allErrors.append("third")
        if type(row4) != NoneType and type(row4) == list:
            if row4[0] == False:
                allErrors.append("fourth")
        if type(row5) != NoneType and type(row5) == list:
            if row5[0] == False:
                allErrors.append("fifth")

        if len(allErrors) > 0:
            errorText = "The "
            for i in range(len(allErrors)):
                if len(allErrors) > 2:
                    if len(errorText) == 4:
                        errorText = errorText + allErrors[i]
                    else:
                        if allErrors[i] == allErrors[-1]:
                            errorText = errorText + ", and " + allErrors[i]
                        else:
                            errorText = errorText + ", " + allErrors[i]
                if 2 >= len(allErrors) > 0:
                    if len(errorText) == 4:
                        errorText = errorText + allErrors[i]
                    else:
                        if allErrors[i] == allErrors[-1]:
                            errorText = errorText + " and " + allErrors[i]
                else:
                    pass
            if len(errorText) > 4:
                errorWindow = Toplevel(root)
                errorWindow.geometry("330x150")
                errorWindow.title("WARNING")
                if len(allErrors) == 1:
                    Label(errorWindow,text =(errorText+" row is a missing a parameter."),font = ("Arial",11)).pack(padx=20,pady=20)
                else:
                    Label(errorWindow,text =(errorText+" rows\n are a missing a parameter."),font = ("Arial",11)).pack(padx=20,pady=20)

                Button(errorWindow,text ="ok",height=1,width=10,command = lambda:closeWindow(errorWindow,False)).pack(padx=20,pady=20)

        else:
            if not checkZero(questions1.get()):
                stored_amnts.append(questions1.get())
                stored_ranges.append([s1_1.get(),s1_2.get()])
                stored_toggles.append(togglebox1.get())
            if not checkZero(questions2.get()):
                stored_amnts.append(questions2.get())
                stored_ranges.append([s2_1.get(),s2_2.get()])
                stored_toggles.append(togglebox2.get())
            if not checkZero(questions3.get()):
                stored_amnts.append(questions3.get())
                stored_ranges.append([s3_1.get(),s3_2.get()])
                stored_toggles.append(togglebox3.get())
            if not checkZero(questions4.get()):
                stored_amnts.append(questions4.get())
                stored_ranges.append([s4_1.get(),s4_2.get()])
                stored_toggles.append(togglebox4.get())
            if not checkZero(questions5.get()):
                stored_amnts.append(questions5.get())
                stored_ranges.append([s5_1.get(),s5_2.get()])
                stored_toggles.append(togglebox5.get())
            for i in range(len(stored_amnts)):
                questionRange = []
                questionsFound = 0
                questions_temp = []
                questionsToggled = 0
                if stored_toggles[i] == "0%":
                    percentage = 0
                elif stored_toggles[i] == "10%":
                    percentage = 0.1
                elif stored_toggles[i] == "20%":
                    percentage = 0.2
                elif stored_toggles[i] == "30%":
                    percentage = 0.3
                elif stored_toggles[i] == "40%":
                    percentage = 0.4
                maxToggled = math.ceil(int(stored_amnts[i])*percentage)
                for chapter in all_chapters:
                    if stored_ranges[i][0] <= chapter <= stored_ranges[i][1]:
                        questionRange.append(chapter)
                for question in all_questions:
                    for questionIndex in questionRange:
                        if question[chapterIndex] == questionIndex:
                            questions_temp.append(question)
                while questionsFound < int(stored_amnts[i]):
                    maybeQuestion = questions_temp[random.randint(0,len(questions_temp)-1)]
                    
                    #checks if question is already in the bank
                    questionIndexes = list(dict.fromkeys(questionIndexes))
                    indexeslength = len(questionIndexes)
                    questionIndexes.append(maybeQuestion[8])
                    questionIndexes = list(dict.fromkeys(questionIndexes))
                    if len(questionIndexes) != indexeslength:
                        if len(maybeQuestion[7]) > 1:
                            if questionsToggled < maxToggled:
                                questionsToggled += 1
                                questionsFound += 1
                                questionsReturned.append(maybeQuestion)
                                questionIndexes.append(maybeQuestion[8])
                        else:
                            questionsFound += 1
                            questionsReturned.append(maybeQuestion)
                            questionIndexes.append(maybeQuestion[8])

            for i in range(len(questionsReturned)):
                values = randomizeQuestion(questionsReturned[i][1],questionsReturned[i][2],questionsReturned[i][3],questionsReturned[i][4],questionsReturned[i][7])
                questionsReturned[i][1] = values[0]
                questionsReturned[i][2] = values[1]
                questionsReturned[i][3] = values[2]
                questionsReturned[i][4] = values[3]
                questionsReturned[i][7] = values[4]
                print(values)
            quiz = Document()
            totalQuestions = 0
            for amnt in stored_amnts:
                totalQuestions += int(amnt)
            test_marks = str(totalQuestions)
            test_question_instruction = "Identify the choice that best completes the statement or answers the question."
            print(COURSE_NAME,TEST_NAME,TEST_WEIGHT,TEST_TIME_LIMIT)
            cover_title = coverTitle(quiz,COURSE_NAME,TEST_NAME,TEST_WEIGHT,TEST_TIME_LIMIT,test_marks)
            quiz.add_page_break()
            instruct1 = quiz.add_paragraph(test_question_instruction)
            instruct1.paragraph_format.first_line_indent = Inches(-0.5)
            instruct1.paragraph_format.first_line_indent.pt
            instructFont1 =instruct1.style.font
            instructFont1.name = "Times New Roman"
            instructFont1.size = Pt(11)
            for i in range(len(questionsReturned)):
                textthing = ("____   ",str(i+1),". ",questionsReturned[i][0])
                question1 = quiz.add_paragraph(textthing)
                question1.paragraph_format.first_line_indent = Inches(-0.5)
                question1.paragraph_format.first_line_indent.pt
                if i == 0:
                    question1.style = quiz.styles.add_style('Style Name', WD_STYLE_TYPE.PARAGRAPH)
                font=question1.style.font
                font.name = 'Times New Roman'
                font.size = Pt(11)
                answerTable = quiz.add_table(rows=2,cols=2)
                colOneCells = answerTable.columns[0].cells
                colTwoCells = answerTable.columns[1].cells
                colOneCells[0].text = "a.   " + questionsReturned[i][1]
                colOneCells[1].text = "b.   " + questionsReturned[i][2]
                colTwoCells[0].text = "c.   " + questionsReturned[i][3]
                colTwoCells[1].text = "d.   " + questionsReturned[i][4]    
            now = datetime.now()
            current_time = now.strftime("__%m-%d-%y-%H.%M.%S")
            TEST_NAME = TEST_NAME.replace(" ","_")
            quizName = COURSE_NAME + "_" + TEST_NAME + alphabet[j]  + current_time +".docx" 
            FILE_NAME = quizName
            if h == "1":
                FIRST_FILE = quizName
            print(FILE_NAME)
            try:
                if not os.path.exists("Quizes"):
                    os.makedirs("Quizes")
            except OSError:
                print ('Error: Creating directory "Quizes". ')
            quiz.save('./Quizes/' + quizName)


            quizAnswers = Document()
            test_question_instruction = "Identify the choice that best completes the statement or answers the question."
            cover_title = coverTitle(quizAnswers,COURSE_NAME,TEST_NAME+" ANSWERS",TEST_WEIGHT,TEST_TIME_LIMIT,test_marks)
            quizAnswers.add_page_break()
            instruct1 = quiz.add_paragraph(test_question_instruction)
            instruct1.paragraph_format.first_line_indent = Inches(-0.5)
            instruct1.paragraph_format.first_line_indent.pt
            instructFont1 =instruct1.style.font
            instructFont1.name = "Times New Roman"
            instructFont1.size = Pt(11)
            for i in range(len(questionsReturned)):
                textthing1 = ()      
                if len(questionsReturned[i][7]) == 1:
                    textthing1 = (questionsReturned[i][7],"___   ",str(i+1),". ",questionsReturned[i][0])
                elif len(questionsReturned[i][7]) == 2:
                    textthing1 = (questionsReturned[i][7],"__   ",str(i+1),". ",questionsReturned[i][0])
                elif len(questionsReturned[i][7]) == 3:
                    textthing1 = (questionsReturned[i][7],"_   ",str(i+1),". ",questionsReturned[i][0])
                # print(questionsReturned[i][7])
                # print(textthing1) 

                question1 = quizAnswers.add_paragraph(textthing1)
                question1.paragraph_format.first_line_indent = Inches(-0.5)
                question1.paragraph_format.first_line_indent.pt
                font=question1.style.font
                font.name = 'Times New Roman'
                font.size = Pt(11)
                answerTable = quizAnswers.add_table(rows=2,cols=2)
                colOneCells = answerTable.columns[0].cells
                colTwoCells = answerTable.columns[1].cells
                colOneCells[0].text = "a.   " + questionsReturned[i][1]
                colOneCells[1].text = "b.   " + questionsReturned[i][2]
                colTwoCells[0].text = "c.   " + questionsReturned[i][3]
                colTwoCells[1].text = "d.   " + questionsReturned[i][4]

            if int(h) == int(COPIES):
                popUp = Toplevel(root)
                Label(popUp,text="Your Quiz has been successfully created.").pack(padx=20,pady=20)
                Button(popUp,text ="ok",height=1,width=10,command = lambda:closeWindow(popUp,True)).pack(padx=20,pady=20)
            quizName1 = COURSE_NAME + "_" + TEST_NAME + alphabet[j] + "_Answers" + current_time +".docx" 
            try:
                if not os.path.exists("Quizes"):
                    os.makedirs("Quizes")
            except OSError:
                print ('Error: Creating directory "Quizes". ')
            quizAnswers.save('./Quizes/' + quizName1)


            
    
    



def moreEntries():
    row1 = checkInsides(questions1,s1_1,s1_2)
    row2 = checkInsides(questions2,s2_1,s2_2)
    row3 = checkInsides(questions3,s3_1,s3_2)
    row4 = checkInsides(questions4,s4_1,s4_2)
    row5 = checkInsides(questions5,s5_1,s5_2)

    allErrors = []
    if type(row1) != NoneType and type(row1) == list:
        if row1[0] == False:
            allErrors.append("first")
    if type(row2) != NoneType and type(row2) == list:
        if row2[0] == False:
            allErrors.append("second")
    if type(row3) != NoneType and type(row3) == list:
        if row3[0] == False:
            allErrors.append("third")
    if type(row4) != NoneType and type(row4) == list:
        if row4[0] == False:
            allErrors.append("fourth")
    if type(row5) != NoneType and type(row5) == list:
        if row5[0] == False:
            allErrors.append("fifth")
    if len(allErrors) > 0:
        errorText = "The "
        for i in range(len(allErrors)):
            if len(allErrors) > 2:
                if len(errorText) == 4:
                    errorText = errorText + allErrors[i]
                else:
                    if allErrors[i] == allErrors[-1]:
                        errorText = errorText + ", and " + allErrors[i]
                    else:
                        errorText = errorText + ", " + allErrors[i]
            if 2 >= len(allErrors) > 0:
                if len(errorText) == 4:
                    errorText = errorText + allErrors[i]
                else:
                    if allErrors[i] == allErrors[-1]:
                        errorText = errorText + " and " + allErrors[i]
            else:
                pass
        if len(errorText) > 4:
            errorWindow = Toplevel(root)
            errorWindow.geometry("330x150")
            errorWindow.title("WARNING")
            if len(allErrors) == 1:
                Label(errorWindow,text =(errorText+" row is a missing a parameter."),font = ("Arial",11)).pack(padx=20,pady=20)
            else:
                Label(errorWindow,text =(errorText+" rows\n are a missing a parameter."),font = ("Arial",11)).pack(padx=20,pady=20)

            Button(errorWindow,text ="ok",height=1,width=10,command = lambda:closeWindow(errorWindow,False)).pack(padx=20,pady=20)
    else:
        if not checkZero(questions1.get()):
            stored_amnts.append(questions1.get())
            stored_ranges.append([s1_1.get(),s1_2.get()])
            stored_toggles.append(togglebox1.get())
            printCheck()
        if not checkZero(questions2.get()):
            stored_amnts.append(questions2.get())
            stored_ranges.append([s2_1.get(),s2_2.get()])
            stored_toggles.append(togglebox2.get())
        if not checkZero(questions3.get()):
            stored_amnts.append(questions3.get())
            stored_ranges.append([s3_1.get(),s3_2.get()])
            stored_toggles.append(togglebox3.get())
        if not checkZero(questions4.get()):
            stored_amnts.append(questions4.get())
            stored_ranges.append([s4_1.get(),s4_2.get()])
            stored_toggles.append(togglebox4.get())
        if not checkZero(questions5.get()):
            stored_amnts.append(questions5.get())
            stored_ranges.append([s5_1.get(),s5_2.get()])
            stored_toggles.append(togglebox5.get())
        clearEntries()
        totalQuestions = 0
        for amnt in stored_amnts:
            totalQuestions += int(amnt)
        valuesText = [("You have currently stored "+ str(totalQuestions)+ " questions. \n\n" )]
        for i in range(len(stored_amnts)):
            if stored_amnts[i] != 0:
                thing = stored_amnts[i] +" questions between " + stored_ranges[i][0] + " and " + stored_ranges[i][1] + "with a maximum of "+ stored_toggles[i] + " of them being multiple choice" + "\n"
                valuesText.append(thing)
        valuesText = "".join(valuesText)

        questions1.focus_set()
        bottomBox.delete("1.0","end")
        bottomBox.insert(INSERT,valuesText)

def clearEntries():
    questions1.delete(0,"end")
    questions2.delete(0,"end")
    questions3.delete(0,"end")
    questions4.delete(0,"end")
    questions5.delete(0,"end")
    s1_1.set("")
    s1_2.set("")
    s2_1.set("")
    s2_2.set("")
    s3_1.set("")
    s3_2.set("")
    s4_1.set("")
    s4_2.set("")
    s5_1.set("")
    s5_2.set("")

def randomizeQuestion(choice1,choice2,choice3,choice4,answers):
    rightChoices = []
    finalChoices = []
    choices = [choice1,choice2,choice3,choice4]
    for answer in answers:
        if answer == "A":
            rightChoices.append(choice1)
        if answer == "B":
            rightChoices.append(choice2)
        if answer == "C":
            rightChoices.append(choice3)
        if answer == "D":
            rightChoices.append(choice4)
    random.shuffle(choices)
    for choice in choices:
        for rightChoice in rightChoices:
            if choice == rightChoice:
                finalChoices.append(choices.index(choice))
    for i in range(len(finalChoices)):
        if finalChoices[i] == 0:
            finalChoices[i] = "A"
        elif finalChoices[i] == 1:
            finalChoices[i] = "B"
        elif finalChoices[i] == 2:
            finalChoices[i] = "C"
        elif finalChoices[i] == 3:
            finalChoices[i] = "D"
    stringthing = ""
    for item in finalChoices:
        stringthing = stringthing + item
    choices.append(stringthing)
    return choices

def root_on_closing():
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        root.destroy()



getCoverInfo()
    




#top labels for all columns
column1 = Label(root,text="No.")
column2 = Label(root,text="Number of Questions")
column3 = Label(root,text="From")
column4 = Label(root,text="To")
column5 = Label(root,text="Limit Multiple \nAnswer Questions")

#Multiple Choice Toggle Boxes
toggleoptions = ["0%","10%","20%","30%","40%"]
toggle1 = StringVar()
toggle2 = StringVar()
toggle3 = StringVar()
toggle4 = StringVar()
toggle5 = StringVar()
togglebox1 = ttk.Combobox(root, textvariable=toggle1, value = toggleoptions,width=12)
togglebox1.current(1)
togglebox2 = ttk.Combobox(root, textvariable=toggle2, value = toggleoptions,width=12)
togglebox2.current(1)
togglebox3 = ttk.Combobox(root, textvariable=toggle3, value = toggleoptions,width=12)
togglebox3.current(1)
togglebox4 = ttk.Combobox(root, textvariable=toggle4, value = toggleoptions,width=12)
togglebox4.current(1)
togglebox5 = ttk.Combobox(root, textvariable=toggle5, value = toggleoptions,width=12)
togglebox5.current(1)

#entry boxes for the amount of questions
questions1 = Entry(root,width = 20)
questions2 = Entry(root,width = 20)
questions3 = Entry(root,width = 20)
questions4 = Entry(root,width = 20)
questions5 = Entry(root,width = 20)

#filler labels
rowLabel1 = Label(root,text="1")
rowLabel2 = Label(root,text="2")
rowLabel3 = Label(root,text="3")
rowLabel4 = Label(root,text="4")
rowLabel5 = Label(root,text="5")

#entry boxes for the categories of the questions
var1_1 = StringVar()
var1_2 = StringVar()

var2_1 = StringVar()
var2_2 = StringVar()

var3_1 = StringVar()
var3_2 = StringVar()

var4_1 = StringVar()
var4_2 = StringVar()

var5_1 = StringVar()
var5_2 = StringVar()

s1_1 = ttk.Combobox(root, textvariable=var1_1, value = all_chapters)
s1_1.bind("<<ComboboxSelected>>", option_reduction1)
s1_2 = ttk.Combobox(root, textvariable=var1_2, value = all_chapters)

s2_1 = ttk.Combobox(root, textvariable=var2_1, value = all_chapters)
s2_1.bind("<<ComboboxSelected>>", option_reduction2)
s2_2 = ttk.Combobox(root, textvariable=var2_2, value = all_chapters)

s3_1 = ttk.Combobox(root, textvariable=var3_1, value = all_chapters)
s3_1.bind("<<ComboboxSelected>>", option_reduction3)
s3_2 = ttk.Combobox(root, textvariable=var3_2, value = all_chapters)

s4_1 = ttk.Combobox(root, textvariable=var4_1, value = all_chapters)
s4_1.bind("<<ComboboxSelected>>", option_reduction4)
s4_2 = ttk.Combobox(root, textvariable=var4_2, value = all_chapters)

s5_1 = ttk.Combobox(root, textvariable=var5_1, value = all_chapters)
s5_1.bind("<<ComboboxSelected>>", option_reduction5)
s5_2 = ttk.Combobox(root, textvariable=var5_2, value = all_chapters)


bottomBox = ScrolledText(root, width=80, height=10,font = ("Arial",12))
bottomBox.insert(INSERT,"You have not saved any entries yet.")
bottomBox.grid(column = 0,row=10, pady = 10, padx = 10,columnspan=5)

submitButton = Button(root, text="Generate PDF",command=makePDF,width=15)
moreEntriesButton = Button(root, text="More Entries",command=moreEntries,width=10)
resetButton = Button(root, text="Clear Entries",command=clearEntries,width=15)
closeButton = Button(root,text="Close",command=lambda:root_on_closing(),width=10)

courseLabel = Label(root,text="Course Name: "ICS4U",anchor="w",font=("Arial",15))
instructLabel = Label(root,text="Pick questions from range of each chapter.",anchor="w",font=("Arial",15))
noteLabel = Label(root,text="Note: Pick from the dropdown list of chapters. Do not input.",anchor="w",font=("Arial",11))

courseLabel.grid(column=0,row=0,padx=40,pady=25,columnspan=5,sticky="W")
instructLabel.grid(column=0,row=1,padx=40,pady=5,columnspan=5,sticky="W")
noteLabel.grid(column=0,row=2,padx=40,pady=10,columnspan=5,sticky="W")

root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=3)
root.grid_columnconfigure(2, weight=3)
root.grid_columnconfigure(3, weight=3)
root.grid_columnconfigure(4, weight=3)



column1.grid(column=0,row=4,padx=20)
column2.grid(column=1,row=4)
column3.grid(column=2,row=4)
column4.grid(column=3,row=4)
column5.grid(column=4,row=4)

questions1.grid(column=1,row=5)
rowLabel1.grid(column=0,row=5)
s1_1.grid(column=2,row=5)
s1_2.grid(column=3,row=5)
togglebox1.grid(column=4,row=5)

questions2.grid(column=1,row=6)
rowLabel2.grid(column=0,row=6)
s2_1.grid(column=2,row=6)
s2_2.grid(column=3,row=6)
togglebox2.grid(column=4,row=6)

questions3.grid(column=1,row=7)
rowLabel3.grid(column=0,row=7)
s3_1.grid(column=2,row=7)
s3_2.grid(column=3,row=7)
togglebox3.grid(column=4,row=7)

questions4.grid(column=1,row=8)
rowLabel4.grid(column=0,row=8)
s4_1.grid(column=2,row=8)
s4_2.grid(column=3,row=8)
togglebox4.grid(column=4,row=8)

questions5.grid(column=1,row=9)
rowLabel5.grid(column=0,row=9)
s5_1.grid(column=2,row=9)
s5_2.grid(column=3,row=9)
togglebox5.grid(column=4,row=9)

closeButton.grid(column=4,row=11,padx=20,pady=20)
submitButton.grid(column=3,row=11,pady=20)
resetButton.grid(column=2,row=11,pady=20)
moreEntriesButton.grid(column=0,row=11,pady=20,columnspan=2)

root.protocol("WM_DELETE_WINDOW", root_on_closing)

root.mainloop()

